"""Processor for Word documents (.docx).

Extracts text from paragraphs and tables, plus embedded images.
Requires ``python-docx``: ``pip install attachments[docx]``
"""

from __future__ import annotations

import logging
from typing import Any

from . import register_processor

log = logging.getLogger("attachments.processors.docx")


def _extract_table_text(table) -> str:
    """Render a docx table as a simple pipe-delimited text table."""
    rows: list[str] = []
    for row in table.rows:
        cells = [cell.text.strip() for cell in row.cells]
        rows.append("| " + " | ".join(cells) + " |")
    return "\n".join(rows)


def docx_processor(data: bytes, **options: Any) -> dict[str, Any]:
    """Convert .docx bytes to an artifact.

    Options:
        filename: Original filename (for metadata).
        images: If ``True``, extract embedded images (default ``False``).
    """
    filename = options.get("filename", "document.docx")
    extract_images = bool(options.get("images") or options.get("render_images"))

    try:
        from docx import Document
    except ImportError:
        return {
            "text": "",
            "images": [],
            "audio": [],
            "video": [],
            "flags": {
                "error": (
                    "Word document processing requires python-docx. "
                    "Install with: pip install attachments[docx]"
                ),
                "filename": filename,
            },
        }

    import io

    try:
        doc = Document(io.BytesIO(data))
    except Exception as e:
        log.warning("failed to open docx %s: %s", filename, e)
        return {
            "text": "",
            "images": [],
            "audio": [],
            "video": [],
            "flags": {
                "error": f"Failed to parse Word document: {e}",
                "filename": filename,
            },
        }

    # --- text extraction (paragraphs + tables in document order) ---
    parts: list[str] = []
    # Walk body elements to preserve paragraph/table ordering
    for element in doc.element.body:
        tag = element.tag.split("}")[-1] if "}" in element.tag else element.tag
        if tag == "p":
            # Find matching paragraph object
            for para in doc.paragraphs:
                if para._element is element:
                    text = para.text.strip()
                    if text:
                        parts.append(text)
                    break
        elif tag == "tbl":
            for table in doc.tables:
                if table._element is element:
                    parts.append(_extract_table_text(table))
                    break

    full_text = "\n\n".join(parts)

    # --- image extraction ---
    images: list[dict[str, Any]] = []
    if extract_images:
        for i, rel in enumerate(doc.part.rels.values()):
            if "image" in rel.reltype:
                try:
                    img_part = rel.target_part
                    img_bytes = img_part.blob
                    ct = img_part.content_type or "image/png"
                    ext = ct.split("/")[-1].split("+")[0]
                    images.append(
                        {
                            "name": f"{filename}-image-{i + 1}.{ext}",
                            "mimetype": ct,
                            "bytes": img_bytes,
                        }
                    )
                except Exception as exc:
                    log.debug("skipping image %d: %s", i, exc)

    return {
        "text": full_text,
        "images": images,
        "audio": [],
        "video": [],
        "flags": {
            "kind": "document",
            "filename": filename,
            "paragraphs": len(doc.paragraphs),
            "tables": len(doc.tables),
            "images_extracted": len(images),
        },
    }


register_processor(".docx", docx_processor)
