"""Tests for the Word (.docx) processor."""

from __future__ import annotations

import io

import pytest

from attachments.deps import check_dep
from attachments.processors import processors

pytestmark = pytest.mark.skipif(
    not check_dep("docx").available,
    reason="python-docx not installed",
)


# ---------------------------------------------------------------------------
# Fixtures
# ---------------------------------------------------------------------------


@pytest.fixture
def simple_docx_bytes():
    """A minimal .docx with two paragraphs."""
    from docx import Document

    doc = Document()
    doc.add_heading("Test Document", level=1)
    doc.add_paragraph("First paragraph with some text.")
    doc.add_paragraph("Second paragraph.")
    buf = io.BytesIO()
    doc.save(buf)
    return buf.getvalue()


@pytest.fixture
def docx_with_table():
    """A .docx containing a table."""
    from docx import Document

    doc = Document()
    doc.add_paragraph("Before table.")
    table = doc.add_table(rows=2, cols=2)
    table.cell(0, 0).text = "Name"
    table.cell(0, 1).text = "Score"
    table.cell(1, 0).text = "Alice"
    table.cell(1, 1).text = "95"
    doc.add_paragraph("After table.")
    buf = io.BytesIO()
    doc.save(buf)
    return buf.getvalue()


@pytest.fixture
def docx_with_image():
    """A .docx containing a tiny embedded PNG image."""
    pytest.importorskip("PIL")
    from docx import Document
    from docx.shared import Inches
    from PIL import Image as PILImage

    # Create a proper 4x4 red PNG via Pillow
    img = PILImage.new("RGB", (4, 4), color="red")
    png_buf = io.BytesIO()
    img.save(png_buf, format="PNG")
    png_buf.seek(0)

    doc = Document()
    doc.add_paragraph("Has image below.")
    doc.add_picture(png_buf, width=Inches(0.5))
    buf = io.BytesIO()
    doc.save(buf)
    return buf.getvalue()


# ---------------------------------------------------------------------------
# Tests
# ---------------------------------------------------------------------------


class TestDocxProcessor:
    def test_registered(self):
        assert ".docx" in processors

    def test_basic_extraction(self, simple_docx_bytes):
        proc = processors[".docx"]
        result = proc(simple_docx_bytes, filename="test.docx")

        assert "Test Document" in result["text"]
        assert "First paragraph" in result["text"]
        assert "Second paragraph" in result["text"]
        assert result["flags"]["kind"] == "document"
        assert result["flags"]["filename"] == "test.docx"

    def test_artifact_structure(self, simple_docx_bytes):
        result = processors[".docx"](simple_docx_bytes)
        for key in ("text", "images", "audio", "video", "flags"):
            assert key in result
        assert isinstance(result["images"], list)

    def test_table_extraction(self, docx_with_table):
        result = processors[".docx"](docx_with_table)

        assert "Before table" in result["text"]
        assert "After table" in result["text"]
        assert "Alice" in result["text"]
        assert "95" in result["text"]
        assert result["flags"]["tables"] >= 1

    def test_image_extraction_off_by_default(self, docx_with_image):
        result = processors[".docx"](docx_with_image)
        assert result["images"] == []

    def test_image_extraction_when_requested(self, docx_with_image):
        result = processors[".docx"](docx_with_image, images=True)
        assert len(result["images"]) >= 1
        img = result["images"][0]
        assert "name" in img
        assert "mimetype" in img
        assert "bytes" in img
        assert len(img["bytes"]) > 0

    def test_corrupt_data_returns_error(self):
        result = processors[".docx"](b"not a docx file")
        assert "error" in result["flags"]
        assert result["text"] == ""

    def test_empty_document(self):
        from docx import Document

        doc = Document()
        buf = io.BytesIO()
        doc.save(buf)
        result = processors[".docx"](buf.getvalue())
        assert isinstance(result["text"], str)
        assert "error" not in result["flags"]


class TestDocxMissingDep:
    @pytest.mark.skipif(
        check_dep("docx").available,
        reason="Test only relevant when python-docx is missing",
    )
    def test_missing_dep_returns_error(self):
        result = processors[".docx"](b"%fake")
        assert "error" in result["flags"]
        assert "python-docx" in result["flags"]["error"]
